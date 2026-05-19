from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "Copia de E_bussiness_corte_3_corregido.docx"
BACKUP_PATH = ROOT / "Copia de E_bussiness_corte_3_corregido.backup.docx"
EVIDENCE_DIR = ROOT / "Evidencias corte 3 Savia Col"
WEB_URL = "https://briyithestrada-90ab7.web.app/"
REPORT_URL = "https://briyithestrada-90ab7.web.app/pages/informe-academico.html"
GA_ID = "G-40E6XHBPJY"


EVIDENCES = [
    ("01", "Mailchimp - flujo de bienvenida", "Prueba el flujo welcome new subscriber configurado para Savia Col.", "01_mailchimp_bienvenida_flujo.jpg"),
    ("02", "Mailchimp - carrito abandonado simulado", "Prueba la automatizacion de carrito abandonado mediante etiqueta Cliente.", "02_mailchimp_carrito_abandonado_simulado.jpg"),
    ("03", "Mailchimp - postcompra y recompra", "Prueba el flujo post-purchase upsell / recompra.", "03_mailchimp_postcompra_recompra.convertido.jpg"),
    ("04", "GA4 - propiedad demo", "Prueba la creacion de la propiedad Savia Col. E-commerce Demo.", "04_ga4_propiedad.convertido.jpg"),
    ("05", "GA4 - flujo web", "Prueba el web stream asociado al portal publicado.", "05_ga4_web_stream.convertido.jpg"),
    ("06", "GA4 - tiempo real", "Prueba que el sitio envia actividad a Analytics.", "06_ga4_realtime.convertido.jpg"),
    ("07", "GA4 - eventos", "Prueba los eventos de e-commerce enviados desde el portal.", "07_ga4_eventos.convertido.jpg"),
    ("09", "GA4 - exploracion de embudo", "Prueba el funnel de abandono de carrito configurado en Exploraciones.", "09_ga4_funnel_abandono.convertido.jpg"),
    ("10", "HBR - pricing dinamico", "Referencia aplicada a precios, promociones y percepcion de confianza.", "10_hbr_pricing_dinamico.jpg"),
    ("11", "HBR - omnicanalidad", "Referencia aplicada a web, WhatsApp, redes, email y entregas.", "11_hbr_omnicanalidad.jpg"),
    ("12", "HBR - pruebas A/B", "Referencia aplicada a mensajes, precios y campanas controladas.", "12_hbr_ab_testing.jpg"),
    ("13", "HBR - matriz estrategica", "Sintesis de los insights HBR usados en el Go-to-Market.", "13_hbr_tabla_estrategica.jpg"),
    ("14", "HubSpot - plantilla pitch", "Resumen ejecutivo, mercado, portafolio y DOFA.", "14_hubspot_plantilla_pitch.jpg"),
    ("15", "HubSpot - finanzas y operaciones", "Complemento del pitch con marketing, finanzas y operacion.", "15_hubspot_finanzas_operaciones.jpg"),
    ("16", "Simulacion IA / CEOSim", "Decisiones, resultados simulados y recomendacion tipo mentor IA.", "16_simulacion_ia_decisiones_resultados.jpg"),
    ("17", "VR tour logistico", "Recorrido visual de almacenamiento, empaque, despacho y postventa.", "17_vr_tour_logistico.jpg"),
    ("18", "Web publicada - documento soporte", "Captura del informe academico dentro del aplicativo web desplegado.", "18_web_publicada_informe.png"),
]


def set_run(run, bold=False, size=None):
    run.bold = bold
    if size:
        run.font.size = Pt(size)


def heading(paragraph, text, size=14):
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run(run, bold=True, size=size)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def normal(paragraph, text, bold_start=None):
    paragraph.text = ""
    if bold_start and text.startswith(bold_start):
        r1 = paragraph.add_run(bold_start)
        set_run(r1, bold=True)
        paragraph.add_run(text[len(bold_start):])
    else:
        paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    return p


def replace_text(doc, replacements):
    for p in doc.paragraphs:
        text = p.text
        if text in replacements:
            p.text = replacements[text]
        else:
            for old, new in replacements.items():
                if old in text:
                    p.text = text.replace(old, new)


def insert_final_status(doc):
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "2. Evaluación de pasarelas de pago internacionales":
            target = p
            break
    if not target:
        return

    p = target.insert_paragraph_before("")
    heading(p, "1.1. Actualizacion final del aplicativo y evidencias")
    items = [
        f"El aplicativo quedo publicado en Firebase Hosting en el enlace principal {WEB_URL} y el documento soporte quedo visible dentro del menu Informe: {REPORT_URL}.",
        "El portal conserva catalogo, fichas de producto, carrito, politicas, medios de pago propuestos, informacion de envio, version bilingue y documento academico navegable.",
        f"Se integro Google Analytics 4 directamente en el codigo del aplicativo con el ID de medicion {GA_ID}. El portal envia eventos de e-commerce para view_item, add_to_cart, begin_checkout y purchase academico/simulado.",
        "El despliegue del sitio se realizo mediante Firebase Hosting con el comando npm run firebase:deploy, dejando el entregable disponible en web para revision del profesor.",
        "Las evidencias finales se organizaron en la carpeta Evidencias corte 3 Savia Col, incluyendo Mailchimp, GA4, HBR, HubSpot, simulacion IA/CEOSim, VR tour y captura de la web publicada.",
    ]
    for item in reversed(items):
        bp = target.insert_paragraph_before(item)
        bp.style = "List Paragraph"
        bp.paragraph_format.left_indent = Inches(0.25)
        bp.paragraph_format.space_after = Pt(3)
    target.insert_paragraph_before("")


def insert_closing_control(doc):
    ref = None
    for p in doc.paragraphs:
        if p.text.strip() == "Referencias":
            ref = p
            break
    if not ref:
        return
    ref.insert_paragraph_before("")
    p = ref.insert_paragraph_before("")
    heading(p, "Control de cierre del entregable")
    paragraphs = [
        "Con los ajustes finales, el informe queda alineado con el aplicativo publicado y con las evidencias guardadas. Los componentes obligatorios del corte quedan cubiertos asi: inventario, costos, postventa, lanzamiento, GA4, HBR, pitch tipo HubSpot y simulacion IA/CEOSim con VR tour logistico.",
        "Para HubSpot no fue necesario conectar un CRM real al aplicativo, porque el requisito se resolvio como plantilla de pitch diligenciada. Para CEOSim o VR tour tampoco se requirio una integracion tecnica dentro de la pagina; se presentaron como simulacion estrategica y evidencia visual de apoyo academico.",
        "El unico punto que debe confirmarse antes de enviar es si el docente exige una captura especifica de purchase marcado como evento clave en GA4 o certificado de Skillshop. La medicion del portal y los eventos principales ya quedaron soportados por codigo y por las capturas de GA4 anexas.",
    ]
    for text in reversed(paragraphs):
        ref.insert_paragraph_before(text)
    ref.insert_paragraph_before("")


def append_evidence_appendix(doc):
    doc.add_page_break()
    heading(doc.add_paragraph(), "Anexo. Evidencias finales del corte 3", 16)
    normal(
        doc.add_paragraph(),
        "Este anexo relaciona las capturas y piezas generadas para sustentar el informe enviado al profesor. Los archivos se encuentran en la carpeta Evidencias corte 3 Savia Col y complementan el documento soporte publicado en la web.",
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Grid Table 1 Light"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["No.", "Evidencia", "Soporte dentro del trabajo", "Archivo"]):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for number, title, support, filename in EVIDENCES:
        row = table.add_row().cells
        row[0].text = number
        row[1].text = title
        row[2].text = support
        row[3].text = filename

    doc.add_paragraph("")
    normal(
        doc.add_paragraph(),
        "Nota: si el profesor solicita evidencia adicional de conversion, se recomienda anexar una captura puntual de GA4 donde purchase aparezca marcado como evento clave. El resto de soportes principales ya queda incluido en este anexo.",
    )

    for idx, (number, title, _support, filename) in enumerate(EVIDENCES, start=1):
        image_path = EVIDENCE_DIR / filename
        if not image_path.exists():
            continue
        doc.add_page_break()
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Evidencia {number}. {title}")
        set_run(r, bold=True, size=11)
        pic_par = doc.add_paragraph()
        pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_par.add_run().add_picture(str(image_path), width=Inches(6.4))


def main():
    if not BACKUP_PATH.exists():
        copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(DOC_PATH)

    replace_text(
        doc,
        {
            "En una implementación real, se utilizaría el ID de medición generado por GA4 y se instalaría en la página mediante Google Tag Manager o mediante el código de Google tag.": f"En el entregable final, el ID de medicion {GA_ID} quedo integrado directamente en el codigo del portal publicado, por medio de Google tag.",
            "6.2.3.3. Instalación simulada mediante Google Tag Manager o código": "6.2.3.3. Instalacion implementada mediante Google tag en el codigo del portal",
            "Para Savia Col., se recomienda usar Google Tag Manager (GTM), porque permite administrar etiquetas y eventos sin modificar directamente el código del sitio cada vez que se quiera medir una nueva acción.": "Para Savia Col., la medicion inicial quedo implementada de forma directa en el codigo del sitio. A futuro, Google Tag Manager puede usarse si se desea administrar nuevas etiquetas sin modificar el codigo cada vez.",
            "La segunda actividad consistió en configurar eventos de comercio electrónico que permitan medir el recorrido del cliente dentro del portal. Para el caso de Savia Col., se proponen cuatro eventos principales:": "La segunda actividad consistio en configurar eventos de comercio electronico que permitan medir el recorrido del cliente dentro del portal. Para el caso de Savia Col., se dejaron definidos cuatro eventos principales:",
            "Los flujos descritos requieren que el portal de Savia Col. registre correctamente los eventos de e-commerce mediante Google Analytics 4, tal como se describe en la sección de analítica. La conexión entre GA4, Mailchimp y el portal permite que las automatizaciones se disparen automáticamente sin intervención manual.": "Los flujos descritos se soportan con los eventos de e-commerce registrados en Google Analytics 4, tal como se describe en la seccion de analitica. En el entregable academico se dejaron simulados en Mailchimp los flujos de bienvenida, carrito abandonado y postcompra, considerando las limitaciones del plan gratuito.",
            "Para una integración completa, se recomiendan los siguientes pasos:": "Para una integracion completa en produccion, los pasos recomendados serian:",
            "Semana 17. Pitch final y plantilla tipo HubSpot": "Semana 17. Pitch final y plantilla tipo HubSpot",
            "Para la Semana 17 se desarrolla el cierre ejecutivo del e-commerce Savia Col. Despensa Saludable S.A.S., utilizando una estructura tipo plantilla HubSpot. Esta sección integra los principales componentes del plan de negocio: productos, mercado, marketing, finanzas, operaciones y pitch final.": "Para la Semana 17 se desarrolla el cierre ejecutivo del e-commerce Savia Col. Despensa Saludable S.A.S., utilizando una estructura tipo plantilla HubSpot. Esta seccion integra los principales componentes del plan de negocio: productos, mercado, marketing, finanzas, operaciones y pitch final. La plantilla quedo soportada con las evidencias 14 y 15 anexas al informe.",
            "Para Savia Col., este elemento puede utilizarse de manera académica y estratégica como un recurso visual que muestre el proceso interno del negocio.": "Para Savia Col., este elemento se utilizo de manera academica y estrategica como un recurso visual que muestra el proceso interno del negocio. La evidencia final corresponde al archivo 17_vr_tour_logistico.jpg.",
            "A partir de las decisiones anteriores, se simula un resultado inicial para los primeros tres meses de operación.": "A partir de las decisiones anteriores, se simula un resultado inicial para los primeros tres meses de operacion. La evidencia final de esta simulacion corresponde al archivo 16_simulacion_ia_decisiones_resultados.jpg.",
        },
    )

    insert_final_status(doc)
    insert_closing_control(doc)
    append_evidence_appendix(doc)

    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
